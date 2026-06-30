from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SRC = BASE / "gpr_detector_system_design.md"
OUT = BASE / "设计文件.docx"
OUT_ASCII = BASE / "gpr_detector_system_design_current.docx"

BODY_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 35, 45)
MUTED = RGBColor(90, 96, 105)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "BFC7D5"


def set_run_font(run, name=BODY_FONT, east_asia=BODY_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_format(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_in=6.5):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def split_table_row(line):
    text = line.strip().strip("|")
    return [part.strip() for part in text.split("|")]


def is_separator_row(line):
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def add_table(doc, rows):
    if not rows:
        return
    header = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table)

    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_para_format(p, after=0, line=1.05)
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=INK, bold=True)

    for row in body:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_para_format(p, after=0, line=1.05)
            if i > 0 and re.fullmatch(r"[+\-]?\d+(\.\d+)?(\s?[-–]\s?\d+(\.\d+)?)?\s?(kg|m|mm|cm)?|约?\s?\d+(\.\d+)?\s?(kg|m|mm|cm)", text):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=9.2, color=INK)

    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_format(header, after=0, line=1)
    run = header.add_run("频域探雷器整机系统设计")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(footer, after=0, line=1)
    run = footer.add_run("设计文件 · 当前硬件尺寸重量版")
    set_run_font(run, size=8.5, color=MUTED)


def add_title_block(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_format(p, before=8, after=4, line=1.05)
    run = p.add_run(title)
    set_run_font(run, size=22, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    set_para_format(p, before=0, after=14, line=1.1)
    run = p.add_run("当前硬件尺寸与重量版 | 2026-06-16")
    set_run_font(run, size=10.5, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, 6.2)
    set_table_borders(meta, "D7DBE2", "4")
    pairs = [
        ("设计对象", "频域体制探雷器整机系统"),
        ("核心约束", "工作长度 1.5-1.7 m；收纳长度 0.7-0.9 m；原目标重量 2.5-3.2 kg"),
        ("当前硬件重量", "已知硬件合计 2.74 kg，结构件与线缆尚未计入"),
        ("本版判断", "推荐第一版结构样机 3.7-4.0 kg；强度验证样机可短期接近 4.2 kg"),
    ]
    for row, (k, v) in zip(meta.rows, pairs):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_FILL)
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        set_para_format(p0, after=0, line=1.05)
        set_para_format(p1, after=0, line=1.05)
        set_run_font(p0.add_run(k), size=9.5, bold=True, color=INK)
        set_run_font(p1.add_run(v), size=9.5, color=INK)
    doc.add_paragraph()


def add_code_block(doc, text):
    lines = text.rstrip("\n").splitlines()
    if not lines:
        return
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    set_table_borders(table, "D7DBE2", "4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    set_para_format(p, after=0, line=1.0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name=MONO_FONT, east_asia=BODY_FONT, size=8.2, color=INK)
    doc.add_paragraph()


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, after=6, line=1.12)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_list_item(doc, text, ordered=False):
    p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    set_para_format(p, after=4, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def build():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)
    set_header_footer(doc)

    first_heading_consumed = False
    in_code = False
    code_lines = []
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            add_code_block(doc, "\n".join(code_lines))
            code_lines = []

    for line in md:
        raw = line.rstrip("\n")

        if raw.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_table()
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(raw)
            continue

        if raw.strip().startswith("|") and raw.strip().endswith("|"):
            if is_separator_row(raw):
                continue
            table_rows.append(split_table_row(raw))
            continue
        flush_table()

        text = raw.strip()
        if not text:
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", text)
        if heading:
            hashes, content = heading.groups()
            if len(hashes) == 1 and not first_heading_consumed:
                add_title_block(doc, content)
                first_heading_consumed = True
            else:
                level = min(len(hashes), 3)
                p = doc.add_heading(content, level=level)
                for run in p.runs:
                    set_run_font(
                        run,
                        size={1: 16, 2: 13, 3: 12}[level],
                        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
                        bold=True,
                    )
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", text)
        if numbered:
            add_list_item(doc, numbered.group(1), ordered=True)
            continue

        bullet = re.match(r"^-\s+(.*)$", text)
        if bullet:
            add_list_item(doc, bullet.group(1), ordered=False)
            continue

        add_body_paragraph(doc, text)

    flush_table()
    flush_code()
    doc.save(OUT)
    shutil.copyfile(OUT, OUT_ASCII)
    print(OUT)
    print(OUT_ASCII)


if __name__ == "__main__":
    build()
